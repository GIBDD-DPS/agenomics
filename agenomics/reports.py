"""
reports.py — форматированные отчёты методологии Agenomics (Markdown и DOCX).

Автор: Dm.Andreyanov
Проект: Prizolov Lab
Версия: 0.4.3

Оборачивает TrustResult/TeamCompatibilityResult в готовые к показу
клиенту отчёты: trust_report()/compatibility_report() — Markdown (без
внешних зависимостей), trust_report_docx() — Word-документ с брендингом
Prizolov Lab (требует python-docx, ставится отдельно: pip install
python-docx — это НЕ входит в основные зависимости пакета, чтобы ядро
методологии оставалось без внешних библиотек).

[v0.4.3] Все три функции принимают language="ru"|"en" — параметр
управляет ЗАГОЛОВКАМИ/ПОДПИСЯМИ самого отчёта. recommendations и
capped_reason внутри result уже локализованы на этапе TrustScorer(
language=...)/CompatibilityScorer(language=...) — если вы хотите
полностью англоязычный отчёт, оба параметра (у Scorer и у функции
отчёта) должны совпадать; report-функции сами это не проверяют.
"""

from .compatibility import TeamCompatibilityResult
from .trust_score import TrustResult

# --- UI-строки самого отчёта (не путать с текстом внутри TrustResult) -------

_UI_STRINGS = {
    "ru": {
        "trust_title": "Trust Score: {agent_id}",
        "score_label": "Score",
        "confidence_label": "Confidence",
        "confidence_of_data": "данных",
        "breakdown_header": "Разбивка по осям",
        "insufficient_flag": " ⚠️ insufficient data",
        "cap_applied": "⚠️ **Потолок применён:**",
        "recommendations_header": "Рекомендации",
        "how_to_prefix": "_Как сделать:_",
        "compat_title": "Compatibility Score команды",
        "avg_compat_label": "Средняя совместимость",
        "weakest_link_label": "Самое слабое звено",
        "all_pairs_header": "Все пары",
        "roles_considered": " (роли учтены)",
        # DOCX-специфичные:
        "docx_report_title": "Trust Score Report",
        "docx_agent_label": "Агент:",
        "docx_overall_score": "ИТОГОВЫЙ БАЛЛ",
        "docx_status": "СТАТУС",
        "docx_confidence": "CONFIDENCE",
        "docx_cap_applied": "⚠️  ПОТОЛОК ПРИМЕНЁН",
        "docx_how_to_label": "Как сделать:  ",
    },
    "en": {
        "trust_title": "Trust Score: {agent_id}",
        "score_label": "Score",
        "confidence_label": "Confidence",
        "confidence_of_data": "data available",
        "breakdown_header": "Breakdown by axis",
        "insufficient_flag": " ⚠️ insufficient data",
        "cap_applied": "⚠️ **Ceiling applied:**",
        "recommendations_header": "Recommendations",
        "how_to_prefix": "_How to:_",
        "compat_title": "Team Compatibility Score",
        "avg_compat_label": "Average compatibility",
        "weakest_link_label": "Weakest link",
        "all_pairs_header": "All pairs",
        "roles_considered": " (roles accounted for)",
        # DOCX-специфичные:
        "docx_report_title": "Trust Score Report",
        "docx_agent_label": "Agent:",
        "docx_overall_score": "OVERALL SCORE",
        "docx_status": "STATUS",
        "docx_confidence": "CONFIDENCE",
        "docx_cap_applied": "⚠️  CEILING APPLIED",
        "docx_how_to_label": "How to:  ",
    },
}


def _ui(language: str, key: str) -> str:
    strings = _UI_STRINGS.get(language, _UI_STRINGS["ru"])
    return strings.get(key, _UI_STRINGS["ru"][key])


# --- Markdown-отчёты (без внешних зависимостей) -----------------------------


def trust_report(result: TrustResult, agent_id: str = "agent", language: str = "ru") -> str:
    lines = [
        f"## {_ui(language, 'trust_title').format(agent_id=agent_id)}",
        "",
        f"**{_ui(language, 'score_label')}:** {result.score}/100 → **{result.label}**",
        f"**{_ui(language, 'confidence_label')}:** {result.confidence} "
        f"({result.confidence_ratio * 100:.0f}% {_ui(language, 'confidence_of_data')})",
        "",
        f"### {_ui(language, 'breakdown_header')}",
    ]
    for axis, value in result.breakdown.items():
        flag = _ui(language, "insufficient_flag") if axis in result.insufficient_axes else ""
        lines.append(f"- {axis}: {value:.0f}/100{flag}")

    if result.capped_reason:
        lines += ["", f"{_ui(language, 'cap_applied')} {result.capped_reason}"]

    if result.recommendations:
        lines += ["", f"### {_ui(language, 'recommendations_header')}"]
        for i, rec in enumerate(result.recommendations, 1):
            lines.append(f"{i}. {rec}")
            # Ищем ось, к которой относится рекомендация, чтобы подставить how_to
            for axis, how_to_text in result.how_to.items():
                if axis in rec:
                    lines.append(f"   {_ui(language, 'how_to_prefix')} {how_to_text}")
                    break

    lines += ["", f"_{result.attribution}_"]
    return "\n".join(lines)


def compatibility_report(result: TeamCompatibilityResult, language: str = "ru") -> str:
    lines = [
        f"## {_ui(language, 'compat_title')}",
        "",
        f"**{_ui(language, 'avg_compat_label')}:** {result.average_score}/100",
        (
            f"**{_ui(language, 'weakest_link_label')}:** {result.weakest_pair.agent_a} ↔ "
            f"{result.weakest_pair.agent_b} ({result.weakest_pair.score}/100)"
        ),
        "",
        f"### {_ui(language, 'all_pairs_header')}",
    ]
    for p in result.pairs:
        role_note = _ui(language, "roles_considered") if p.complementary_roles else ""
        lines.append(f"- {p.agent_a} ↔ {p.agent_b}: {p.score}/100{role_note}")
        if p.capped_reason:
            lines.append(f"  ⚠️ {p.capped_reason}")

    attribution = result.pairs[0].attribution if result.pairs else ""
    lines += ["", f"_{attribution}_"]
    return "\n".join(lines)


# --- DOCX-отчёт (требует python-docx, опциональная зависимость) -------------

_NAVY = "1B2A4A"
_NAVY_LIGHT = "2C4270"
_ORANGE = "E65100"
_ORANGE_BG = "FFF3E0"
_GREEN = "2E7D32"
_RED = "C62828"
_GRAY = "6B7280"
_LIGHT = "F4F5F7"
_WHITE = "FFFFFF"

_STATUS_COLORS = {"Trusted": _GREEN, "Conditional": _ORANGE, "High Risk": _RED}
_AXIS_LABELS = {
    "transparency": "Transparency",
    "bias_control": "Bias Control",
    "data_safety": "Data Safety",
    "predictability": "Predictability",
    "accountability": "Accountability",
}
_CONTENT_WIDTH_IN = 7.5  # US Letter, поля 0.5" с каждой стороны


def _require_python_docx():
    try:
        import docx  # noqa: F401
    except ImportError as e:
        raise ImportError(
            "trust_report_docx() требует пакет 'python-docx', который не "
            "входит в основные зависимости agenomics (ядро методологии "
            "остаётся без внешних библиотек). Установите отдельно:\n"
            "    pip install python-docx"
        ) from e


def trust_report_docx(
    result: TrustResult, agent_id: str, output_path: str, language: str = "ru"
) -> str:
    """
    Генерирует брендированный Word-отчёт (.docx) по Trust Score и
    сохраняет его по указанному пути.

    language управляет заголовками/подписями САМОГО ОТЧЁТА ("ru"|"en").
    Текст recommendations/capped_reason/how_to внутри result уже
    локализован на этапе TrustScorer(language=...) — для полностью
    англоязычного отчёта передайте language="en" в оба места.

    Требует python-docx (pip install python-docx) — устанавливается
    отдельно, так как ядро agenomics не имеет внешних зависимостей.

    Возвращает output_path.
    """
    _require_python_docx()
    ui = lambda key: _ui(language, key)  # noqa: E731 — локальный алиас для краткости

    from docx import Document
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def _rgb(hex_color):
        return RGBColor.from_string(hex_color)

    def _set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _remove_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tcPr.append(borders)

    def _set_row_height(row, points):
        row.height = Pt(points)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def _set_cantsplit(row):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))

    def _set_table_fixed_layout(table):
        # Без этого LibreOffice/Word может проигнорировать заданную ширину
        # ячеек и распределить колонки поровну — критично для прогресс-баров.
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table._tbl.tblPr.append(layout)

    def _set_col_widths(table, widths_in):
        table.autofit = False
        _set_table_fixed_layout(table)
        for i, w in enumerate(widths_in):
            table.columns[i].width = Inches(w)
        for row in table.rows:
            for cell, w in zip(row.cells, widths_in):
                cell.width = Inches(w)

    def _set_font_all_scripts(run, name="Calibri"):
        # Без явного указания шрифта для latin/eastAsia/cs LibreOffice может
        # не применить bold/italic к кириллическому тексту.
        run.font.name = name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), name)

    def _add_run(paragraph, text, size=11, bold=False, italic=False, color=None):
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = _rgb(color)
        _set_font_all_scripts(run)
        return run

    def _new_table(doc, rows, cols, widths_in):
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_col_widths(table, widths_in)
        return table

    def _progress_bar(doc, score, bar_color):
        filled_frac = max(0.0, min(1.0, score / 100))
        filled_w = round(_CONTENT_WIDTH_IN * filled_frac, 3)
        empty_w = round(_CONTENT_WIDTH_IN - filled_w, 3)

        if empty_w <= 0.01:
            table = _new_table(doc, 1, 1, [_CONTENT_WIDTH_IN])
            cells = [table.rows[0].cells[0]]
            colors = [bar_color]
        else:
            table = _new_table(doc, 1, 2, [filled_w, empty_w])
            cells = table.rows[0].cells
            colors = [bar_color, "E2E4E9"]

        _set_row_height(table.rows[0], 9)
        _set_cantsplit(table.rows[0])
        for cell, color in zip(cells, colors):
            _set_cell_bg(cell, color)
            _remove_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    def _axis_row(doc, label, score, flagged):
        table = _new_table(doc, 1, 2, [_CONTENT_WIDTH_IN * 0.64, _CONTENT_WIDTH_IN * 0.36])
        _set_cantsplit(table.rows[0])
        left, right = table.rows[0].cells
        _remove_cell_borders(left)
        _remove_cell_borders(right)
        _add_run(left.paragraphs[0], label, size=10.5, bold=True)
        p_right = right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p_right, f"{score:.0f}/100", size=10.5, bold=True, color=_RED if flagged else "222222")
        if flagged:
            _add_run(p_right, "  ⚠", color=_RED)

    def _recommendation_card(doc, num, title, how_to, accent_color, how_to_label):
        table = _new_table(doc, 1, 2, [0.08, _CONTENT_WIDTH_IN - 0.08])
        _set_cantsplit(table.rows[0])
        accent_cell, body_cell = table.rows[0].cells
        _set_cell_bg(accent_cell, accent_color)
        _remove_cell_borders(accent_cell)
        _set_cell_bg(body_cell, _LIGHT)
        _remove_cell_borders(body_cell)
        body_cell.margin_left = Inches(0.12)
        body_cell.margin_right = Inches(0.12)
        body_cell.margin_top = Inches(0.08)
        body_cell.margin_bottom = Inches(0.08)
        _add_run(body_cell.paragraphs[0], f"{num}. {title}", size=11, bold=True)
        p_how = body_cell.add_paragraph()
        p_how.paragraph_format.space_before = Pt(3)
        _add_run(p_how, how_to_label, size=9.5, bold=True, color=_GRAY)
        _add_run(p_how, how_to, size=9.5, color="374151")

    def _colored_band(doc, text_lines, bg_color):
        table = _new_table(doc, 1, 1, [_CONTENT_WIDTH_IN])
        cell = table.rows[0].cells[0]
        _set_cell_bg(cell, bg_color)
        _remove_cell_borders(cell)
        cell.margin_top = Inches(0.12)
        cell.margin_bottom = Inches(0.12)
        cell.margin_left = Inches(0.18)
        cell.margin_right = Inches(0.18)
        first = True
        for line in text_lines:
            text, size, bold, color = line[:4]
            italic = line[4] if len(line) > 4 else False
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            _add_run(p, text, size=size, bold=bold, italic=italic, color=color)

    # --- Сборка документа ----------------------------------------------

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin = section.bottom_margin = Inches(0.5)

    status_color = _STATUS_COLORS.get(result.label, _GRAY)

    # Шапка: два столбца в одной цветной плашке
    header_table = _new_table(doc, 1, 2, [_CONTENT_WIDTH_IN * 0.66, _CONTENT_WIDTH_IN * 0.34])
    _set_cantsplit(header_table.rows[0])
    left_cell, right_cell = header_table.rows[0].cells
    for c in (left_cell, right_cell):
        _set_cell_bg(c, _NAVY)
        _remove_cell_borders(c)
        c.margin_top = Inches(0.2)
        c.margin_bottom = Inches(0.2)
    left_cell.margin_left = Inches(0.22)
    right_cell.margin_right = Inches(0.22)

    _add_run(left_cell.paragraphs[0], "🧬 AGENOMICS", size=10, bold=True, color="9FB3D9")
    p2 = left_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    _add_run(p2, ui("docx_report_title"), size=20, bold=True, color=_WHITE)
    p3 = left_cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    _add_run(p3, ui("docx_agent_label") + "  ", size=10, color="9FB3D9")
    _add_run(p3, agent_id, size=10, bold=True, color=_WHITE)

    pr1 = right_cell.paragraphs[0]
    pr1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(pr1, "PRIZOLOV LAB", size=13, bold=True, color=_WHITE)
    pr2 = right_cell.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr2.paragraph_format.space_before = Pt(2)
    _add_run(pr2, "Dm.Andreyanov", size=9, color="9FB3D9")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Плашка score / статус / confidence
    stats_table = _new_table(doc, 1, 3, [_CONTENT_WIDTH_IN / 3] * 3)
    _set_cantsplit(stats_table.rows[0])
    c_score, c_status, c_conf = stats_table.rows[0].cells
    for c in (c_score, c_status, c_conf):
        _set_cell_bg(c, _LIGHT)
        _remove_cell_borders(c)
        c.margin_top = Inches(0.15)
        c.margin_bottom = Inches(0.15)
        c.margin_left = Inches(0.12)
        c.margin_right = Inches(0.12)

    _add_run(c_score.paragraphs[0], ui("docx_overall_score"), size=8, bold=True, color=_GRAY)
    p2 = c_score.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    _add_run(p2, f"{result.score:.0f}", size=28, bold=True, color=status_color)
    _add_run(p2, "/100", size=11, color=_GRAY)

    _add_run(c_status.paragraphs[0], ui("docx_status"), size=8, bold=True, color=_GRAY)
    p2 = c_status.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    _add_run(p2, result.label, size=15, bold=True, color=status_color)

    _add_run(c_conf.paragraphs[0], ui("docx_confidence"), size=8, bold=True, color=_GRAY)
    p2 = c_conf.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    _add_run(p2, result.confidence, size=15, bold=True, color=_GREEN)
    _add_run(p2, f"  ({result.confidence_ratio * 100:.0f}%)", size=9, color=_GRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Разбивка по осям
    h = doc.add_heading(level=2)
    _add_run(h, ui("breakdown_header"), size=14, bold=True, color=_NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for axis, value in result.breakdown.items():
        label = _AXIS_LABELS.get(axis, axis)
        flagged = axis in result.how_to
        _axis_row(doc, label, value, flagged)
        _progress_bar(doc, value, _RED if flagged else _NAVY_LIGHT)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Потолок (если применён)
    if result.capped_reason:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        _colored_band(
            doc,
            [
                (ui("docx_cap_applied"), 10.5, True, _ORANGE),
                (result.capped_reason, 9.5, False, "5A3A00"),
            ],
            _ORANGE_BG,
        )

    # Рекомендации
    if result.recommendations:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        h2 = doc.add_heading(level=2)
        _add_run(h2, ui("recommendations_header"), size=14, bold=True, color=_NAVY)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        accent_cycle = [_RED, _ORANGE, _ORANGE, _GRAY, _GRAY]
        for i, rec in enumerate(result.recommendations, 1):
            how_to_text = ""
            for axis, text in result.how_to.items():
                if axis in rec:
                    how_to_text = text
                    break
            accent = accent_cycle[min(i - 1, len(accent_cycle) - 1)]
            _recommendation_card(doc, i, rec, how_to_text, accent, ui("docx_how_to_label"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Подвал
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    _colored_band(doc, [(result.attribution, 8.5, False, "C7D2E8", True)], _NAVY)

    doc.save(output_path)
    return output_path
