"""
test_i18n.py — тесты мультиязычности (v0.4.3): language в TrustScorer,
CompatibilityScorer, trust_report(), compatibility_report(), trust_report_docx().

Автор: Dm.Andreyanov
Проект: Prizolov Lab
Версия: 0.4.3
"""

import os
import tempfile

from agenomics import (
    AgentGenome, TrustScorer, CompatibilityScorer,
    trust_report, compatibility_report, trust_report_docx,
    SUPPORTED_LANGUAGES,
)

try:
    import docx as _docx_check  # noqa: F401
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _assert_raises(exc_type, fn):
    try:
        fn()
    except exc_type:
        return
    raise AssertionError(f"Ожидалось исключение {exc_type.__name__}")


def test_supported_languages_contains_ru_and_en():
    assert "ru" in SUPPORTED_LANGUAGES
    assert "en" in SUPPORTED_LANGUAGES


def test_trust_scorer_unknown_language_raises():
    _assert_raises(ValueError, lambda: TrustScorer(language="fr"))


def test_compatibility_scorer_unknown_language_raises():
    _assert_raises(ValueError, lambda: CompatibilityScorer(language="fr"))


def test_trust_scorer_language_does_not_affect_score_value():
    """Язык влияет только на текст, не на сами числа — иначе это баг."""
    genome = AgentGenome(
        id="x", domain="finance", autonomy="autonomous",
        transparency=75, bias_control=80, data_safety=85,
        drift_rate=0.1, has_ledger=False,
    )
    ru = TrustScorer(language="ru").score(genome)
    en = TrustScorer(language="en").score(genome)
    assert ru.score == en.score
    assert ru.label == en.label  # label ("Trusted" и т.д.) уже на английском в обоих случаях
    assert ru.confidence == en.confidence


def test_trust_scorer_capped_reason_is_translated():
    genome = AgentGenome(
        id="x", domain="finance", autonomy="autonomous",
        transparency=75, bias_control=80, data_safety=85,
        drift_rate=0.1, has_ledger=False,
    )
    ru = TrustScorer(language="ru").score(genome)
    en = TrustScorer(language="en").score(genome)
    assert "Autonomous-агент" in ru.capped_reason
    assert "An Autonomous agent" in en.capped_reason
    assert ru.capped_reason != en.capped_reason


def test_trust_scorer_recommendations_and_how_to_translated():
    genome = AgentGenome(
        id="needs-help", domain="support", autonomy="autonomous",
        transparency=85, bias_control=72, data_safety=90,
        drift_rate=0.35, has_ledger=False,
    )
    en = TrustScorer(language="en").score(genome)
    assert any("Improve accountability" in r for r in en.recommendations)
    assert "Enable a decision log" in en.how_to["accountability"]


def test_compatibility_scorer_capped_reason_translated():
    a = AgentGenome(id="strict", bias_control=95)
    b = AgentGenome(id="loose", bias_control=40)
    ru = CompatibilityScorer(language="ru").score_pair(a, b)
    en = CompatibilityScorer(language="en").score_pair(a, b)
    assert "Этическое расхождение" in ru.capped_reason
    assert "Ethical divergence" in en.capped_reason
    assert ru.score == en.score  # числа не меняются от языка


def test_trust_report_markdown_headers_translated():
    genome = AgentGenome(id="x", bias_control=50, transparency=50, data_safety=50, drift_rate=0.5, has_ledger=False)
    result = TrustScorer(language="en").score(genome)
    report = trust_report(result, agent_id="x", language="en")
    assert "Breakdown by axis" in report
    assert "Разбивка по осям" not in report


def test_trust_report_markdown_default_is_russian():
    genome = AgentGenome(id="x", bias_control=50, transparency=50, data_safety=50, drift_rate=0.5, has_ledger=False)
    result = TrustScorer().score(genome)  # default language="ru"
    report = trust_report(result, agent_id="x")  # default language="ru"
    assert "Разбивка по осям" in report


def test_compatibility_report_translated():
    a = AgentGenome(id="a", bias_control=80, risk_tolerance=50, social_style=50)
    b = AgentGenome(id="b", bias_control=82, risk_tolerance=50, social_style=50)
    result = CompatibilityScorer(language="en").score_team([a, b])
    report = compatibility_report(result, language="en")
    assert "Team Compatibility Score" in report
    assert "Weakest link" in report


def test_docx_report_translated_headers():
    if not HAS_DOCX:
        return
    genome = AgentGenome(
        id="support-bot-v2", domain="support", autonomy="autonomous",
        transparency=85, bias_control=72, data_safety=90,
        drift_rate=0.35, has_ledger=False,
    )
    result = TrustScorer(language="en").score(genome)
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "report.docx")
        trust_report_docx(result, "support-bot-v2", path, language="en")

        from docx import Document
        doc = Document(path)
        all_text = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        full_text = " ".join(all_text)
        assert "OVERALL SCORE" in full_text
        assert "Recommendations" in full_text
        assert "CEILING APPLIED" in full_text
        assert "Разбивка" not in full_text
